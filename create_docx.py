import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in dxa (1 pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_props in kwargs.items():
        node = OxmlElement(f'w:{border_name}')
        for key, val in border_props.items():
            node.set(qn(f'w:{key}'), str(val))
        tcBorders.append(node)
    tcPr.append(tcBorders)

def set_paragraph_bidi(p):
    """Set paragraph to Right-To-Left (RTL)"""
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_run_rtl(run):
    """Set run to RTL"""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    # Set complex script font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), 'Amiri')
    rFonts.set(qn('w:ascii'), 'Segoe UI')
    rPr.append(rFonts)

doc = docx.Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Base Style setup
normal_style = doc.styles['Normal']
font = normal_style.font
font.name = 'Segoe UI'
font.size = Pt(14)

# 1. Title
p_title = doc.add_paragraph()
set_paragraph_bidi(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(10)
p_title.paragraph_format.space_after = Pt(25)

run_title = p_title.add_run("وكـــــالـــــة")
run_title.font.name = 'Amiri'
run_title.font.size = Pt(26)
run_title.bold = True
set_run_rtl(run_title)

# 2. Paragraph 1
p1 = doc.add_paragraph()
set_paragraph_bidi(p1)
p1.paragraph_format.line_spacing = 1.6
p1.paragraph_format.space_after = Pt(12)

r = p1.add_run("أنا الموقع أسفله، السيد: "); set_run_rtl(r)
r = p1.add_run("أبو القمح حميد"); r.bold = True; set_run_rtl(r)
r = p1.add_run("، الحامل للبطاقة الوطنية للتعريف رقم: "); set_run_rtl(r)
r = p1.add_run("N78311"); r.bold = True; set_run_rtl(r)
r = p1.add_run("، والساكن بـ: "); set_run_rtl(r)
r = p1.add_run("[عنوان الأب]"); r.bold = True; set_run_rtl(r)
r = p1.add_run("."); set_run_rtl(r)

# 3. Paragraph 2
p2 = doc.add_paragraph()
set_paragraph_bidi(p2)
p2.paragraph_format.line_spacing = 1.6
p2.paragraph_format.space_after = Pt(12)

r = p2.add_run("أفوض بموجب هذه الوكالة ابني السيد: "); set_run_rtl(r)
r = p2.add_run("أبو القمح صلاح الدين"); r.bold = True; set_run_rtl(r)
r = p2.add_run("، الحامل للبطاقة الوطنية للتعريف رقم: "); set_run_rtl(r)
r = p2.add_run("N476536"); r.bold = True; set_run_rtl(r)
r = p2.add_run("، والساكن بـ: "); set_run_rtl(r)
r = p2.add_run("تجزئة الرونق، الصويرة 23"); r.bold = True; set_run_rtl(r)
r = p2.add_run("."); set_run_rtl(r)

# 4. Paragraph 3
p3 = doc.add_paragraph()
set_paragraph_bidi(p3)
p3.paragraph_format.line_spacing = 1.6
p3.paragraph_format.space_after = Pt(15)

r = p3.add_run("وذلك لينوب عني ويقوم مقامي في جميع الإجراءات الإدارية اللازمة لدى مصالح "); set_run_rtl(r)
r = p3.add_run("الوكالة الوطنية للسلامة الطرقية (NARSA)"); r.bold = True; set_run_rtl(r)
r = p3.add_run("، المتعلقة بطلب واستخراج نظير من البطاقة الرمادية الخاصة بمركبتي، وذلك بسبب ضياعها."); set_run_rtl(r)

# 5. Vehicle Table Box
table_v = doc.add_table(rows=3, cols=2)
table_v.alignment = WD_TABLE_ALIGNMENT.CENTER

# Style vehicle box container cell border
for row in table_v.rows:
    for cell in row.cells:
        set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

pv_head = table_v.cell(0, 0).paragraphs[0]
set_paragraph_bidi(pv_head)
r = pv_head.add_run("معلومات المركبة:")
r.bold = True
r.font.size = Pt(15)
set_run_rtl(r)

# Merge header across 2 cols
table_v.cell(0, 0).merge(table_v.cell(0, 1))

# Row 1: Label and Value
p_lbl1 = table_v.cell(1, 0).paragraphs[0]
set_paragraph_bidi(p_lbl1)
r = p_lbl1.add_run("• العلامة التجارية:")
r.bold = True
set_run_rtl(r)

p_val1 = table_v.cell(1, 1).paragraphs[0]
set_paragraph_bidi(p_val1)
r = p_val1.add_run("NASH")
r.bold = True
set_run_rtl(r)

# Row 2: Label and Value
p_lbl2 = table_v.cell(2, 0).paragraphs[0]
set_paragraph_bidi(p_lbl2)
r = p_lbl2.add_run("• الرقم الترتيبي:")
r.bold = True
set_run_rtl(r)

p_val2 = table_v.cell(2, 1).paragraphs[0]
set_paragraph_bidi(p_val2)
r = p_val2.add_run("32-001346")
r.bold = True
set_run_rtl(r)

# 6. Paragraph 4
p4 = doc.add_paragraph()
set_paragraph_bidi(p4)
p4.paragraph_format.line_spacing = 1.6
p4.paragraph_format.space_before = Pt(15)
p4.paragraph_format.space_after = Pt(25)

r = p4.add_run("وقد سلمت هذه الوكالة للمعني بالأمر قصد القيام بالإجراءات المذكورة أعلاه، ولتكون له حجة فيما يخص ذلك."); set_run_rtl(r)

# 7. Date line
p_date = doc.add_paragraph()
set_paragraph_bidi(p_date)
p_date.paragraph_format.space_after = Pt(25)
r = p_date.add_run("حرر بالصويرة في: ")
r.bold = True
set_run_rtl(r)
r = p_date.add_run("09/08/2026")
r.bold = True
set_run_rtl(r)

# 8. Signatures Table (2 columns: Legalization Frame & Signature)
sig_table = doc.add_table(rows=1, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_leg = sig_table.cell(0, 0)
cell_sig = sig_table.cell(0, 1)

set_cell_border(cell_leg, 
                top={"sz": 4, "val": "dashed", "color": "888888"},
                bottom={"sz": 4, "val": "dashed", "color": "888888"},
                left={"sz": 4, "val": "dashed", "color": "888888"},
                right={"sz": 4, "val": "dashed", "color": "888888"})

p_leg = cell_leg.paragraphs[0]
set_paragraph_bidi(p_leg)
p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_leg.add_run("إطار خاص بتصحيح الإمضاء\n(Légalisation de Signature)")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(100, 100, 100)
set_run_rtl(r)

p_sig = cell_sig.paragraphs[0]
set_paragraph_bidi(p_sig)
p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sig.add_run("إمضاء الموكل (الأب):\n\n\n....................................................")
r.bold = True
set_run_rtl(r)

# 9. Footnote
p_foot = doc.add_paragraph()
set_paragraph_bidi(p_foot)
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(30)
r = p_foot.add_run("(يتم توقيع الوكالة أمام موظف تصحيح الإمضاء)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(100, 100, 100)
set_run_rtl(r)

# Save files
path_workspace = 'c:/Users/Dark/Desktop/Surf_School_Essaouira/procuration.docx'
path_desktop = 'c:/Users/Dark/Desktop/وكالة_نظير_البطاقة_الرمادية.docx'

doc.save(path_workspace)
doc.save(path_desktop)
print(f"Successfully generated DOCX files:\n - {path_workspace}\n - {path_desktop}")
